"""Import the completed 2026 Orioles DOCX archive into Jekyll posts.

The source documents remain the editorial originals. This script only writes the
known generated post filenames listed below and deliberately leaves the July 27
draft unpublished until it is complete.
"""

from __future__ import annotations

import html
import re
from datetime import date, datetime
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "articles" / "Orioles"
POSTS_DIR = ROOT / "_posts"
PUBLISH_THROUGH = date(2026, 7, 25)

BRANDON_SOURCE = "Brandon Young’s Breakout.docx"
BRANDON_SLUG = "brandon-young-breakout"

WEEKLY_RECAPS = {
    "O_s Weekly Recap 4_6_26.docx": date(2026, 4, 6),
    "O_s Weekly Recap 4_13_26.docx": date(2026, 4, 13),
    "O_s Weekly Recap 4_20_26.docx": date(2026, 4, 20),
    "O_s Weekly Recap 4_27_26.docx": date(2026, 4, 27),
    "O_s Weekly Recap 5_4_26.docx": date(2026, 5, 4),
    "O_s Weekly Recap 5_11_26.docx": date(2026, 5, 11),
    "O_s Weekly Recap 5_18_26.docx": date(2026, 5, 18),
    "O_s Weekly Recap 5_25_26.docx": date(2026, 5, 25),
    "O_s Weekly Recap 6_1_26.docx": date(2026, 6, 1),
    "O_s Weekly Recap 6_8_26.docx": date(2026, 6, 8),
    "O_s Weekly Recap 6_15_26.docx": date(2026, 6, 15),
    "O_s Weekly Recap 6_22_26.docx": date(2026, 6, 22),
    "O_s Weekly Recap 7_13_26.docx": date(2026, 7, 13),
    "O_s Weekly Recap 7_27_26.docx": date(2026, 7, 27),
}

SECTION_LABELS = {"Player of the Week", "Down on the Farm", "Question of the Week"}


def yaml_quote(value: str) -> str:
    return '"' + value.replace("\\", "\\\\").replace('"', '\\"') + '"'


def normalize_text(value: str) -> str:
    value = value.replace("\u00a0", " ").replace("\u2011", "-")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def markdown_run(text: str, *, bold: bool, italic: bool) -> str:
    text = text.replace("\u00a0", " ").replace("\u2011", "-")
    if not text:
        return ""
    leading = text[: len(text) - len(text.lstrip())]
    trailing = text[len(text.rstrip()) :]
    text = text.strip()
    if not text:
        return leading
    text = text.replace("\\", "\\\\")
    text = re.sub(r"([*_`])", r"\\\1", text)
    if bold and italic:
        text = f"***{text}***"
    elif bold:
        text = f"**{text}**"
    elif italic:
        text = f"*{text}*"
    return leading + text + trailing


def paragraph_markdown(paragraph, *, skip_chars: int = 0) -> str:
    chunks: list[str] = []
    remaining_skip = skip_chars
    for run in paragraph.runs:
        run_text = run.text
        if remaining_skip:
            consumed = min(remaining_skip, len(run_text))
            run_text = run_text[consumed:]
            remaining_skip -= consumed
        rendered = markdown_run(run_text, bold=bool(run.bold), italic=bool(run.italic))
        if rendered:
            chunks.append(rendered)
    rendered = "".join(chunks)
    rendered = re.sub(r"\s+", " ", rendered)
    return rendered.strip()


def plain_paragraphs(document: Document) -> list[str]:
    return [normalize_text(paragraph.text) for paragraph in document.paragraphs]


def excerpt_from(document: Document, *, skip: int = 1, max_chars: int = 240) -> str:
    candidates = [text for text in plain_paragraphs(document)[skip:] if len(text) >= 80]
    if not candidates:
        return ""
    excerpt = candidates[0]
    if len(excerpt) <= max_chars:
        return excerpt
    shortened = excerpt[: max_chars - 1].rsplit(" ", 1)[0]
    return shortened.rstrip(".,;:") + "…"


def front_matter(fields: list[tuple[str, str]]) -> str:
    rows = ["---"]
    rows.extend(f"{key}: {value}" for key, value in fields)
    rows.extend(["---", ""])
    return "\n".join(rows)


def body_from_document(document: Document, *, skip_first: bool = True) -> str:
    body: list[str] = []
    paragraphs = document.paragraphs[1:] if skip_first else document.paragraphs
    for paragraph in paragraphs:
        text = paragraph_markdown(paragraph)
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        plain = normalize_text(paragraph.text)
        if style_name.startswith("Heading"):
            try:
                level = max(2, min(4, int(style_name.split()[-1])))
            except (TypeError, ValueError):
                level = 2
            text = f"{'#' * level} {plain}"
        body.append(text)
    return "\n\n".join(body).strip() + "\n"


def split_scorecard(summary: str) -> list[tuple[str, str]]:
    entries: list[tuple[str, str]] = []
    for line in summary.splitlines():
        line = normalize_text(line)
        if not line or ":" not in line:
            continue
        label, value = line.split(":", 1)
        entries.append((normalize_text(label), normalize_text(value)))
    return entries


def weekly_body(document: Document) -> str:
    nonempty = [paragraph for paragraph in document.paragraphs if normalize_text(paragraph.text)]
    if len(nonempty) < 3:
        raise ValueError("Weekly recap does not contain enough content")

    summary_entries = split_scorecard(nonempty[1].text)
    body: list[str] = []
    if summary_entries:
        body.append('<div class="orioles-recap-scorecard" aria-label="Weekly recap summary">')
        for label, value in summary_entries:
            body.append(
                "  <span><small>"
                + html.escape(label)
                + "</small><strong>"
                + html.escape(value)
                + "</strong></span>"
            )
        body.append("</div>")

    for paragraph in nonempty[2:]:
        plain = normalize_text(paragraph.text)
        rendered = paragraph_markdown(paragraph)
        if plain in SECTION_LABELS:
            rendered = f"## {plain}"
        else:
            leading_label = next(
                (label for label in SECTION_LABELS if plain.startswith(label + " ")),
                None,
            )
            if leading_label:
                remainder = paragraph_markdown(paragraph, skip_chars=len(leading_label)).strip()
                body.append(f"## {leading_label}")
                if remainder:
                    body.append(remainder)
                continue
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            rendered = f"## {plain}"
        body.append(rendered)

    return "\n\n".join(body).strip() + "\n"


def import_brandon() -> Path:
    source = SOURCE_DIR / BRANDON_SOURCE
    document = Document(source)
    title = normalize_text(document.paragraphs[0].text)
    excerpt = (
        "Brandon Young’s improvement is real—but the ERA still overstates it. "
        "A closer look at the pitch changes behind his breakout and what Baltimore can expect next."
    )
    output = POSTS_DIR / f"2026-07-25-{BRANDON_SLUG}.md"
    content = front_matter(
        [
            ("layout", "post"),
            ("title", yaml_quote(title)),
            ("date", "2026-07-25 08:00:00 -0400"),
            ("categories", "[orioles]"),
            ("section_theme", "orioles"),
            ("permalink", f"/orioles/{BRANDON_SLUG}/"),
            ("author", yaml_quote("Randy Morgan")),
            ("excerpt", yaml_quote(excerpt)),
            (
                "hero_image",
                yaml_quote(
                    "https://img.mlbstatic.com/mlb-images/image/upload/"
                    "t_2x1/t_w1536/mlb/axrx1rdj9xkzyspakzeu.jpg"
                ),
            ),
            ("hero_wide", "true"),
            ("hero_credit", yaml_quote("Photo: MLB.com — Brandon Young at Houston, July 19, 2026.")),
            (
                "hero_credit_url",
                yaml_quote(
                    "https://www.mlb.com/news/"
                    "brandon-young-7-solid-innings-orioles-extend-win-streak-7-games"
                ),
            ),
            ("suppress_article_prompts", "true"),
            (
                "article_stats",
                '{ title: "2025 vs. 2026", headers: ["2025", "2026"], '
                'note: "Statistics through July 19, 2026", '
                'source_label: "MLB / Baseball Savant", '
                'source_url: "https://baseballsavant.mlb.com/savant-player/brandon-young-687064", '
                'rows: ['
                '{ name: "Record", values: ["1–7", "8–2"] }, '
                '{ name: "ERA", values: ["6.24", "3.25"] }, '
                '{ name: "xERA", values: ["4.27", "4.22"] }, '
                '{ name: "xFIP", values: ["4.52", "4.53"] }, '
                '{ name: "WHIP", values: ["1.54", "1.31"] }, '
                '{ name: "K%", values: ["18.4", "19.0"] }, '
                '{ name: "BB%", values: ["8.6", "8.2"] }, '
                '{ name: "Barrel%", values: ["10.2", "6.1"] }, '
                '{ name: "Slider use", values: ["8.9", "14.5"] }, '
                '{ name: "Slider whiff%", values: ["19.4", "43.0"] }, '
                '{ name: "Chase%", values: ["27.8", "34.8"] }] }',
            ),
            ("orioles_feature", "true"),
            ("source_docx", yaml_quote(f"articles/Orioles/{BRANDON_SOURCE}")),
            ("generated_by", yaml_quote("scripts/import_orioles_articles.py")),
        ]
    ) + body_from_document(document)
    output.write_text(content, encoding="utf-8", newline="\n")
    return output


def import_weekly(source_name: str, publication_date: date) -> Path:
    source = SOURCE_DIR / source_name
    document = Document(source)
    title = f"Orioles Weekly Recap: {publication_date.strftime('%B')} {publication_date.day}, {publication_date.year}"
    slug = f"weekly-recap-{publication_date.isoformat()}"
    excerpt = excerpt_from(document, skip=2)
    output = POSTS_DIR / f"{publication_date.isoformat()}-orioles-{slug}.md"
    content = front_matter(
        [
            ("layout", "post"),
            ("title", yaml_quote(title)),
            ("date", f"{publication_date.isoformat()} 08:00:00 -0400"),
            ("categories", "[orioles]"),
            ("section_theme", "orioles"),
            ("permalink", f"/orioles/{slug}/"),
            ("author", yaml_quote("Randy Morgan")),
            ("excerpt", yaml_quote(excerpt)),
            ("series", yaml_quote("Orioles Weekly Recap")),
            ("suppress_article_prompts", "true"),
            ("source_docx", yaml_quote(f"articles/Orioles/{source_name}")),
            ("generated_by", yaml_quote("scripts/import_orioles_articles.py")),
        ]
    ) + weekly_body(document)
    output.write_text(content, encoding="utf-8", newline="\n")
    return output


def main() -> None:
    missing = [name for name in [BRANDON_SOURCE, *WEEKLY_RECAPS] if not (SOURCE_DIR / name).exists()]
    if missing:
        raise FileNotFoundError("Missing source documents: " + ", ".join(missing))

    written = [import_brandon()]
    skipped: list[str] = []
    for source_name, publication_date in WEEKLY_RECAPS.items():
        if publication_date > PUBLISH_THROUGH:
            skipped.append(source_name)
            continue
        written.append(import_weekly(source_name, publication_date))

    print(f"Imported {len(written)} Orioles posts:")
    for path in written:
        print(f"  {path.relative_to(ROOT)}")
    if skipped:
        print("Left unpublished because the dated source is not complete/current:")
        for name in skipped:
            print(f"  articles/Orioles/{name}")


if __name__ == "__main__":
    main()
